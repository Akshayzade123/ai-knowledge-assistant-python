"""
Embedding Service - Handles document ingestion and embedding.

SOLID Principles Applied:
- Single Responsibility (S): Only handles document processing and embedding
- Dependency Inversion (D): Depends on interfaces for LLM, vector store, and repositories
- Open/Closed (O): Can support new document types without modifying core logic
"""

import logging
from pathlib import Path
from typing import Any

from app.interfaces.database import IDocumentRepository
from app.interfaces.llm import IEmbeddingProvider
from app.interfaces.vector_store import IVectorStore, VectorDocument

logger = logging.getLogger(__name__)


class DocumentChunk:
    """Represents a chunk of a document."""

    def __init__(self, content: str, metadata: dict[str, Any], chunk_index: int):
        self.content = content
        self.metadata = metadata
        self.chunk_index = chunk_index


class EmbedService:
    """
    Service for document ingestion and embedding.

    This service orchestrates the process of:
    1. Loading and chunking documents
    2. Generating embeddings
    3. Storing in vector database
    4. Recording metadata in relational database
    """

    def __init__(
        self,
        embedding_provider: IEmbeddingProvider,
        vector_store: IVectorStore,
        document_repository: IDocumentRepository,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        """
        Initialize embedding service.

        Args:
            embedding_provider: Provider for generating embeddings
            vector_store: Vector database for storing embeddings
            document_repository: Repository for document metadata
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks
        """
        self.embedding_provider = embedding_provider
        self.vector_store = vector_store
        self.document_repository = document_repository
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        logger.info("EmbedService initialized")

    def chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Input text to chunk

        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < text_length:
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)

                if break_point > self.chunk_size // 2:
                    chunk = chunk[: break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return chunks

    async def ingest_document(
        self,
        file_path: str,
        title: str,
        uploaded_by: int,
        department: str | None,
        access_level: str,
        collection_name: str = "Documents",
    ) -> int:
        """
        Ingest a document: chunk, embed, and store.

        Args:
            file_path: Path to document file
            title: Document title
            uploaded_by: User ID who uploaded
            department: Department (for access control)
            access_level: Access level (public, department, restricted)
            collection_name: Vector store collection name

        Returns:
            Document ID from database

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type not supported
        """
        try:
            logger.info(f"Ingesting document: {title}")

            # Load document content
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            file_type = path.suffix.lower()
            content = await self._load_document(file_path, file_type)

            # Chunk document
            chunks = self.chunk_text(content)
            logger.info(f"Created {len(chunks)} chunks from document")

            # Generate embeddings for all chunks
            embedding_results = await self.embedding_provider.embed_batch(chunks)

            # Prepare vector documents
            vector_docs = []
            for i, (chunk, embedding_result) in enumerate(
                zip(chunks, embedding_results, strict=False)
            ):
                vector_doc = VectorDocument(
                    id=f"{title}_{i}",
                    content=chunk,
                    embedding=embedding_result.embedding,
                    metadata={
                        "title": title,
                        "file_path": file_path,
                        "department": department or "",
                        "access_level": access_level,
                        "uploaded_by": uploaded_by,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                    },
                )
                vector_docs.append(vector_doc)

            # Store in vector database
            vector_ids = await self.vector_store.add_documents(vector_docs, collection_name)

            # Store metadata in relational database
            document = await self.document_repository.create_document(
                title=title,
                file_path=file_path,
                file_type=file_type,
                uploaded_by=uploaded_by,
                department=department,
                access_level=access_level,
                vector_store_id=vector_ids[0] if vector_ids else None,
            )

            logger.info(f"Successfully ingested document {title} with ID {document.id}")
            return document.id

        except Exception as e:
            logger.error(f"Error ingesting document: {e}")
            raise

    async def _load_document(self, file_path: str, file_type: str) -> str:
        """
        Load document content based on file type.

        Args:
            file_path: Path to file
            file_type: File extension

        Returns:
            Document text content
        """
        if file_type in [".txt", ".md"]:
            with open(file_path, encoding="utf-8") as f:
                return f.read()

        elif file_type == ".pdf":
            # For production, use PyPDF2 or pdfplumber
            # Simplified for demonstration
            try:
                import PyPDF2

                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    return text
            except ImportError:
                logger.warning("PyPDF2 not installed, treating as text file")
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    return f.read()

        elif file_type in [".doc", ".docx"]:
            # For production, use python-docx
            try:
                from docx import Document

                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                logger.warning("python-docx not installed, treating as text file")
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    return f.read()

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def delete_document(self, document_id: int, collection_name: str = "Documents") -> bool:
        """
        Delete a document from both vector store and database.

        Args:
            document_id: Document ID from database
            collection_name: Vector store collection name

        Returns:
            True if successful
        """
        try:
            # Get document metadata
            document = await self.document_repository.get_document_by_id(document_id)
            if not document:
                logger.warning(f"Document {document_id} not found")
                return False

            # Delete from vector store
            if document.vector_store_id:
                await self.vector_store.delete_document(document.vector_store_id, collection_name)

            # Delete from database
            await self.document_repository.delete_document(document_id)

            logger.info(f"Deleted document {document_id}")
            return True

        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False

    async def get_accessible_documents(
        self, user_role: str, user_department: str | None
    ) -> list[dict[str, Any]]:
        """
        Get list of documents accessible to user.

        Args:
            user_role: User's role
            user_department: User's department

        Returns:
            List of document metadata
        """
        try:
            documents = await self.document_repository.get_accessible_documents(
                user_role, user_department
            )

            return [
                {
                    "id": doc.id,
                    "title": doc.title,
                    "file_type": doc.file_type,
                    "department": doc.department,
                    "access_level": doc.access_level,
                    "created_at": doc.created_at.isoformat(),
                }
                for doc in documents
            ]

        except Exception as e:
            logger.error(f"Error getting accessible documents: {e}")
            return []
